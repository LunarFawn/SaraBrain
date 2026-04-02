"""Convert arxiv_preprint.md to a formatted Word document."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = Path(__file__).parent / "arxiv_preprint.md"
DST = Path(__file__).parent / "arxiv_preprint.docx"

doc = Document()

# -- styles --
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    sname = f"Heading {level}"
    s = doc.styles[sname]
    s.font.name = "Times New Roman"
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.font.size = Pt({1: 16, 2: 14, 3: 12}[level])

code_style_name = "CodeBlock"
if code_style_name not in [s.name for s in doc.styles]:
    cs = doc.styles.add_style(code_style_name, 1)  # paragraph
    cs.font.name = "Consolas"
    cs.font.size = Pt(9)
    cs.paragraph_format.space_before = Pt(4)
    cs.paragraph_format.space_after = Pt(4)
    cs.paragraph_format.left_indent = Inches(0.3)


def add_paragraph(text, bold=False, italic=False, style_name=None):
    p = doc.add_paragraph(style=style_name)
    # Handle inline formatting: **bold** and *italic* and `code`
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            run = p.add_run(part)
    if bold:
        for run in p.runs:
            run.bold = True
    if italic:
        for run in p.runs:
            run.italic = True
    return p


def add_table(header_row, data_rows):
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = h.strip()
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = "Times New Roman"
    # data
    for ri, row in enumerate(data_rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val.strip()
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "Times New Roman"
    doc.add_paragraph()  # spacer


lines = SRC.read_text(encoding="utf-8").splitlines()
i = 0
in_code = False
code_buf = []
in_table = False
table_header = []
table_rows = []

while i < len(lines):
    line = lines[i]

    # --- code blocks ---
    if line.strip().startswith("```"):
        if in_code:
            # end code block
            for cl in code_buf:
                doc.add_paragraph(cl, style=code_style_name)
            code_buf = []
            in_code = False
        else:
            # flush any pending table
            if in_table:
                add_table(table_header, table_rows)
                in_table = False
                table_header = []
                table_rows = []
            in_code = True
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # --- tables ---
    if "|" in line and line.strip().startswith("|"):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # check if next line is separator
        if not in_table:
            # peek for separator
            if i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip()):
                in_table = True
                table_header = cells
                i += 2  # skip header + separator
                continue
            else:
                # not a real table, treat as text
                pass
        else:
            table_rows.append(cells)
            # check if next line is still table
            if i + 1 >= len(lines) or "|" not in lines[i + 1] or not lines[i + 1].strip().startswith("|"):
                add_table(table_header, table_rows)
                in_table = False
                table_header = []
                table_rows = []
            i += 1
            continue

    # flush pending table if we hit non-table line
    if in_table:
        add_table(table_header, table_rows)
        in_table = False
        table_header = []
        table_rows = []

    # --- headings ---
    if line.startswith("# ") and not line.startswith("## "):
        # Title
        p = doc.add_heading(line[2:].strip(), level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue
    if line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=1)
        i += 1
        continue
    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=2)
        i += 1
        continue

    # --- horizontal rules ---
    if line.strip() == "---":
        i += 1
        continue

    # --- blockquotes ---
    if line.strip().startswith("> "):
        text = line.strip()[2:]
        p = add_paragraph(text, italic=True)
        p.paragraph_format.left_indent = Inches(0.4)
        i += 1
        continue

    # --- bullet points ---
    if re.match(r'^(\s*)[-*]\s', line):
        text = re.sub(r'^(\s*)[-*]\s', '', line)
        p = add_paragraph(text)
        p.paragraph_format.left_indent = Inches(0.3)
        i += 1
        continue

    # --- numbered lists ---
    if re.match(r'^\d+\.\s', line):
        text = re.sub(r'^\d+\.\s', '', line)
        p = add_paragraph(text)
        p.paragraph_format.left_indent = Inches(0.3)
        i += 1
        continue

    # --- empty lines ---
    if not line.strip():
        i += 1
        continue

    # --- normal paragraph ---
    add_paragraph(line)
    i += 1

# flush
if in_table:
    add_table(table_header, table_rows)

doc.save(str(DST))
print(f"Saved: {DST}")
